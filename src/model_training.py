# import os
# import pickle
# from sklearn.ensemble import RandomForestClassifier
# from sklearn.model_selection import train_test_split
# from sklearn.metrics import accuracy_score

# # Train the model
# def train_model(X, y):
#     # Split the dataset
#     X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
#     model = RandomForestClassifier()
#     model.fit(X_train, y_train)
    
#     # Evaluate the model
#     y_pred = model.predict(X_test)
#     accuracy = accuracy_score(y_test, y_pred)
#     print(f"Model accuracy: {accuracy:.2f}")
    
#     # Ensure the 'models' directory exists
#     if not os.path.exists('models'):
#         os.makedirs('models')

#     # Save the model
#     model_path = 'models/resume_ranker.pkl'
#     with open(model_path, 'wb') as model_file:
#         pickle.dump(model, model_file)
#     print(f"Model saved to {model_path}")

# # Load the trained model
# def load_model(filepath):
#     # Check if the directory and file exist
#     if not os.path.exists(filepath):
#         raise FileNotFoundError(f"The model file '{filepath}' does not exist.")
    
#     with open(filepath, 'rb') as model_file:
#         model = pickle.load(model_file)
#     return model

# # Example usage:
# # Assuming X and y are defined with your data
# # train_model(X, y)
# # model = load_model('models/resume_ranker.pkl')

# import os
# import pandas as pd
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity
# import PyPDF2
# from docx import Document

# def read_file_as_text(file_path):
#     _, ext = os.path.splitext(file_path)
#     if ext == '.txt':
#         with open(file_path, 'r', encoding='utf-8') as f:
#             return f.read()
#     elif ext == '.pdf':
#         with open(file_path, 'rb') as f:
#             reader = PyPDF2.PdfReader(f)
#             text = []
#             for page in reader.pages:
#                 text.append(page.extract_text())
#             return '\n'.join(text)
#     elif ext == '.docx':
#         doc = Document(file_path)
#         return '\n'.join([para.text for para in doc.paragraphs])
#     return ""

# def rank_resumes(resume_path, jd_path, output_csv='ranking_results.csv'):
#     resume_text = read_file_as_text(resume_path)
#     jd_text = read_file_as_text(jd_path)

#     # Ensure both texts are not empty
#     if not resume_text or not jd_text:
#         raise ValueError("Resume or Job Description is empty.")

#     # Vectorize the texts
#     vectorizer = TfidfVectorizer()
#     vectors = vectorizer.fit_transform([resume_text, jd_text])

#     # Calculate cosine similarity
#     score = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]

#     # Create a DataFrame to store the results
#     results = pd.DataFrame({
#         'Resume File': [os.path.basename(resume_path)],
#         'Job Description File': [os.path.basename(jd_path)],
#         'Score': [score]
#     })

#     # Append the results to the CSV file
#     if not os.path.isfile(output_csv):
#         results.to_csv(output_csv, index=False)  # Create a new file if it doesn't exist
#     else:
#         results.to_csv(output_csv, mode='a', header=False, index=False)  # Append to the existing file

#     return score

# # Example usage
# # score = rank_resumes('path/to/resume.pdf', 'path/to/job_description.txt')

import os
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2
from docx import Document

def read_file_as_text(file_path):
    _, ext = os.path.splitext(file_path)
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == '.pdf':
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = []
            for page in reader.pages:
                text.append(page.extract_text())
            return '\n'.join(text)
    elif ext == '.docx':
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    return ""

def rank_resumes(resume_paths, jd_path, output_csv='ranking_results.csv'):
    # Read the job description
    jd_text = read_file_as_text(jd_path)

    # Ensure the job description is not empty
    if not jd_text:
        raise ValueError("Job Description is empty.")

    # Initialize a list to store results
    results = []

    # Vectorizer for TF-IDF
    vectorizer = TfidfVectorizer()

    for resume_path in resume_paths:
        resume_text = read_file_as_text(resume_path)

        # Ensure the resume is not empty
        if not resume_text:
            continue  # Skip empty resumes

        # Vectorize the texts
        vectors = vectorizer.fit_transform([resume_text, jd_text])

        # Calculate cosine similarity
        score = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]

        # Store the result
        results.append({
            'Resume File': os.path.basename(resume_path),
            'Job Description File': os.path.basename(jd_path),
            'Score': score
        })

    # Create a DataFrame from the results
    results_df = pd.DataFrame(results)

    # Append the results to the CSV file
    if not os.path.isfile(output_csv):
        results_df.to_csv(output_csv, index=False)  # Create a new file if it doesn't exist
    else:
        results_df.to_csv(output_csv, mode='a', header=False, index=False)  # Append to the existing file

    return results_df  # Return DataFrame for further use if needed

# Example usage
#scores = rank_resumes(['path/to/resume1.pdf', 'path/to/resume2.docx'], 'path/to/job_description.txt')
